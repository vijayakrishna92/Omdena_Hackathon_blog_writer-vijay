import streamlit as st
from crewai import Crew
from agents import ContentAgents
from tasks import Tasks
from io import BytesIO
from docx import Document
import json

class ContentCrew:
    def __init__(self, topic):
        self.topic = topic
        self.blog_content = None  # Initialize blog_content to store the generated content

    def run(self):
        agents = ContentAgents()  # Create an instance of ContentAgents
        writer = agents.writer_agent(self.topic)

        tasks = Tasks()  # Ensure this is correctly instantiated
        write_task = tasks.tasker(writer, self.topic)

        crew = Crew(
            agents=[writer],
            tasks=[write_task],
            verbose=True
        )
        
        result = crew.kickoff()  # Execute the crew task
        st.write(result)  # Display the entire result for debugging

        # Accessing the task output
        if result.tasks_output and isinstance(result.tasks_output, list):
            task_output = result.tasks_output[0]  # Get the first task output
            self.blog_content = task_output.raw  # Use raw output as blog content

            # Display additional task output information
            st.write(f"Task Description: {task_output.description}")
            st.write(f"Task Summary: {task_output.summary}")
            st.write(f"Raw Output: {task_output.raw}")
            if task_output.json_dict:
                st.write(f"JSON Output: {json.dumps(task_output.json_dict, indent=2)}")
            if task_output.pydantic:
                st.write(f"Pydantic Output: {task_output.pydantic}")
        else:
            self.blog_content = "No content available."

        st.write(self.blog_content)  # Display the blog content

    def download_blog_as_word(self):
        if self.blog_content:  # Ensure there is content to download
            # Create a Word document
            doc = Document()
            doc.add_heading('Blog Post', level=1)
            doc.add_paragraph(self.blog_content)

            # Save it to a BytesIO object
            bio = BytesIO()
            doc.save(bio)
            bio.seek(0)

            # Download the document
            st.download_button(
                label="Download",
                data=bio,
                file_name='blog_post.docx',
                mime='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
        else:
            st.write("No blog content available for download.")

# Streamlit UI
st.title('Blog Writing Assistant')
st.write('-------------------------------')
topic = st.text_input("Enter your blog topic:")
content_crew = None  # Initialize content_crew variable

if st.button('Generate Blog'):
    if topic:
        content_crew = ContentCrew(topic)  # Create an instance of ContentCrew
        content_crew.run()
    else:
        st.write("Please enter a valid topic.")

# Conditional download button to prevent page reload
if content_crew and content_crew.blog_content:
    content_crew.download_blog_as_word()
